# -*- coding: utf-8 -*-
"""
Generador de la memoria técnica de Rage Tracker en formato .docx editable.

Produce un documento Word con:
- Portada (proyecto, autor, titulación, fecha).
- Índice automático con hipervínculos (campo TOC de Word, se actualiza con F9).
- Estilos de título (Título 1 / Título 2) para navegación correcta.
- Numeración de páginas en el pie.
- Bloques de código con formato monoespaciado y sombreado.
- Huecos "[FOTO: ...]" donde el autor insertará capturas reales.

Uso: python scripts/generar_memoria.py
Salida: Rage_Tracker_Memoria.docx en la raíz del proyecto.
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# --------------------------------------------------------------------------- #
#  Utilidades de bajo nivel (campos, sombreado, numeración)                   #
# --------------------------------------------------------------------------- #
def _set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_field(paragraph, instr_text):
    """Inserta un campo de Word (begin/instrText/separate/end)."""
    run = paragraph.add_run()
    fldBegin = OxmlElement("w:fldChar")
    fldBegin.set(qn("w:fldCharType"), "begin")
    run._r.append(fldBegin)

    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    run2._r.append(instr)

    run3 = paragraph.add_run()
    fldSep = OxmlElement("w:fldChar")
    fldSep.set(qn("w:fldCharType"), "separate")
    run3._r.append(fldSep)

    run4 = paragraph.add_run("Actualiza este campo con F9")
    run5 = paragraph.add_run()
    fldEnd = OxmlElement("w:fldChar")
    fldEnd.set(qn("w:fldCharType"), "end")
    run5._r.append(fldEnd)


def add_page_number_footer(doc):
    """Pie con 'Página X de Y' centrado."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Página ")
    _add_field(p, "PAGE")
    p.add_run(" de ")
    _add_field(p, "NUMPAGES")
    for r in p.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x6A, 0x74, 0x88)


def add_toc(doc):
    """Inserta un índice automático de Word (niveles 1-2) con hipervínculos."""
    p = doc.add_paragraph()
    _add_field(
        p,
        'TOC \\o "1-2" \\h \\z \\u',  # \h = hipervínculos, niveles 1 a 2
    )


# --------------------------------------------------------------------------- #
#  Estilo base del documento                                                  #
# --------------------------------------------------------------------------- #
def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    # Color de acento para los títulos (cian oscuro coherente con la app)
    for name, size, color in (
        ("Heading 1", 18, RGBColor(0x12, 0x6E, 0x82)),
        ("Heading 2", 14, RGBColor(0x1C, 0x7A, 0x6E)),
    ):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(14)
        st.paragraph_format.space_after = Pt(6)

    # Estilo de código monoespaciado
    if "CodeBlock" not in [s.name for s in doc.styles]:
        from docx.enum.style import WD_STYLE_TYPE
        code = doc.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        code.font.name = "Consolas"
        code.font.size = Pt(9.5)
        code.paragraph_format.space_after = Pt(2)
        code.paragraph_format.space_before = Pt(2)
        code.paragraph_format.line_spacing = 1.0


# --------------------------------------------------------------------------- #
#  Bloques de contenido reutilizables                                         #
# --------------------------------------------------------------------------- #
def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def para(doc, text):
    p = doc.add_paragraph(text)
    return p


def bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def numbered(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def code_block(doc, code, lang=""):
    """Tabla de una celda con sombreado gris y texto monoespaciado."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _set_cell_bg(cell, "F2F3F5")
    # Borde sutil
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "D0D4DA")
        borders.append(el)
    tblPr.append(borders)

    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
    first = True
    for line in code.rstrip("\n").split("\n"):
        p = cell.add_paragraph(style="CodeBlock")
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    doc.add_paragraph()  # respiro tras el bloque


def photo(doc, caption):
    """Hueco para una captura que insertará el autor."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ FOTO: {caption} ]")
    run.bold = True
    run.font.color.rgb = RGBColor(0xB0, 0x30, 0x30)
    run.font.size = Pt(11)
    box = doc.add_paragraph()
    box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = box.add_run("________________________________________________\n"
                     "(espacio reservado para la imagen)\n"
                     "________________________________________________")
    r2.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    r2.font.size = Pt(9)
    doc.add_paragraph()


def table_simple(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = htext
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()


# --------------------------------------------------------------------------- #
#  Portada e índice                                                           #
# --------------------------------------------------------------------------- #
def cover_page(doc):
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("RAGE TRACKER")
    r.bold = True
    r.font.size = Pt(40)
    r.font.color.rgb = RGBColor(0x12, 0x6E, 0x82)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Sistema multi-sensor de detección y análisis de\n"
                     "emociones, gritos e insultos en tiempo real durante\n"
                     "sesiones de videojuegos")
    rs.italic = True
    rs.font.size = Pt(13)
    rs.font.color.rgb = RGBColor(0x55, 0x5B, 0x66)

    for _ in range(6):
        doc.add_paragraph()

    def field(label, value, bold_value=True):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = p.add_run(f"{label}: ")
        rl.font.size = Pt(13)
        rl.font.color.rgb = RGBColor(0x6A, 0x74, 0x88)
        rv = p.add_run(value)
        rv.font.size = Pt(13)
        rv.bold = bold_value

    field("Memoria técnica del proyecto", "Documentación completa", bold_value=False)
    doc.add_paragraph()
    field("Autor", "Anthony Ramos de León")
    field("Titulación", "Grado de Especialización en Python")
    field("Fecha", date.today().strftime("%d/%m/%Y"))

    doc.add_page_break()


# --------------------------------------------------------------------------- #
#  Construcción del documento                                                 #
# --------------------------------------------------------------------------- #
def build():
    doc = Document()
    setup_styles(doc)
    add_page_number_footer(doc)

    # ---- Portada ----
    cover_page(doc)

    # ---- Índice ----
    idx = doc.add_paragraph()
    rr = idx.add_run("Índice")
    rr.bold = True
    rr.font.size = Pt(20)
    rr.font.color.rgb = RGBColor(0x12, 0x6E, 0x82)
    nota = doc.add_paragraph()
    rn = nota.add_run("Para generar el índice: haz clic sobre él y pulsa F9 → "
                      "«Actualizar toda la tabla». Los títulos ya usan estilos "
                      "de Word, así que los enlaces funcionan automáticamente.")
    rn.italic = True
    rn.font.size = Pt(9)
    rn.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    add_toc(doc)
    doc.add_page_break()

    # =====================================================================
    # 1. INTRODUCCIÓN
    # =====================================================================
    h1(doc, "1. Introducción")

    h2(doc, "1.1 Descripción del proyecto")
    para(doc, "Rage Tracker es una aplicación de escritorio que combina visión por "
              "computadora y análisis de audio en tiempo real para medir la respuesta "
              "emocional de un jugador mientras juega. Durante una partida, el sistema "
              "observa tres señales independientes y complementarias: las emociones "
              "faciales captadas por la webcam, los gritos detectados a partir del "
              "volumen del micrófono y los insultos reconocidos mediante un motor de "
              "voz a texto en español. Toda esa información se consolida en sesiones "
              "que se almacenan localmente y se visualizan en un dashboard web "
              "interactivo.")
    para(doc, "El nombre del proyecto resume su propósito: rastrear el «rage» (la "
              "frustración) que provocan los videojuegos. La aplicación no busca "
              "juzgar al jugador, sino ofrecerle datos objetivos sobre cómo reacciona "
              "ante distintos títulos y situaciones, todo ello sin que ningún dato "
              "salga de su ordenador.")

    h2(doc, "1.2 Problema que resuelve")
    para(doc, "Los jugadores rara vez son conscientes del impacto emocional real que "
              "tienen ciertos juegos sobre ellos. La percepción que uno tiene de su "
              "propia frustración es subjetiva y poco fiable: tendemos a recordar los "
              "picos puntuales y a olvidar la tensión sostenida. No existe una "
              "herramienta sencilla, local y respetuosa con la privacidad que mida de "
              "forma objetiva esa carga emocional sesión a sesión.")
    para(doc, "Rage Tracker cubre ese vacío: cuantifica la frustración (gritos, "
              "expresiones de enfado e insultos) y la satisfacción (sonrisas, rachas "
              "de felicidad) de cada partida, las asocia a un juego concreto y permite "
              "compararlas a lo largo del tiempo.")

    h2(doc, "1.3 Objetivo principal")
    para(doc, "Diseñar y construir un sistema multi-sensor capaz de detectar y "
              "registrar, en tiempo real y de forma totalmente local, las reacciones "
              "emocionales de un jugador durante sus sesiones de juego, presentando los "
              "resultados de manera clara y accionable.")

    h2(doc, "1.4 Objetivos específicos")
    bullets(doc, [
        "Detectar emociones faciales básicas (feliz / enfadado) mediante visión por "
        "computadora robusta a distintas caras, distancias e iluminaciones.",
        "Detectar gritos a partir del nivel de volumen del micrófono con un umbral "
        "configurable y ajustable en caliente.",
        "Detectar insultos en español mediante reconocimiento de voz offline, sin "
        "mostrar ni almacenar nunca la transcripción.",
        "Permitir activar cualquier combinación de los tres sensores antes de cada "
        "sesión.",
        "Persistir las sesiones en un formato portable y legible, con migración "
        "automática de esquema entre versiones.",
        "Ofrecer un dashboard web que resuma estadísticas globales, por juego y por "
        "sesión.",
        "Garantizar el procesamiento 100 % local, sin conexión a internet ni envío "
        "de datos a terceros.",
    ])

    h2(doc, "1.5 Alcance del proyecto")
    para(doc, "El proyecto abarca la aplicación de escritorio completa: captura de "
              "vídeo y audio, los tres detectores, la interfaz nativa (launcher), la "
              "calibración asistida de cámara, la persistencia en CSV, el servidor del "
              "dashboard y el propio dashboard web. También incluye el empaquetado como "
              "ejecutable de Windows para usuarios sin Python instalado.")
    para(doc, "Quedan fuera del alcance actual: el análisis en la nube, las cuentas de "
              "usuario, la versión móvil y el reconocimiento de un repertorio emocional "
              "más amplio (tristeza, sorpresa, etc.), todo ello contemplado como línea "
              "de mejora futura.")

    # =====================================================================
    # 2. ANÁLISIS DE NECESIDADES
    # =====================================================================
    h1(doc, "2. Análisis de necesidades")

    h2(doc, "2.1 Situación actual")
    para(doc, "El mundo del videojuego mueve a millones de jugadores que pasan horas "
              "frente a títulos competitivos o de alta dificultad. Existe una creciente "
              "conciencia sobre la salud mental y el bienestar emocional asociados al "
              "ocio digital, pero las herramientas disponibles para medir ese impacto "
              "son escasas, intrusivas o dependen de servicios en la nube que recopilan "
              "datos personales sensibles (vídeo y audio del usuario en su propia casa).")

    h2(doc, "2.2 Problemas detectados")
    bullets(doc, [
        "La autoevaluación emocional es subjetiva y poco fiable.",
        "Las soluciones existentes suelen requerir hardware específico (pulseras, "
        "sensores biométricos) o suscripciones.",
        "Muchas alternativas envían datos a servidores externos, lo que genera "
        "desconfianza al tratarse de cámara y micrófono.",
        "No hay una herramienta que vincule la reacción emocional a un juego concreto "
        "para poder compararlos.",
        "El usuario medio no quiere configurar entornos complejos: necesita algo que "
        "funcione al abrirlo.",
    ])

    h2(doc, "2.3 Necesidades de los usuarios")
    bullets(doc, [
        "Saber qué juegos les generan más frustración y cuáles más disfrute.",
        "Disponer de datos objetivos y no de impresiones.",
        "Tener garantías de privacidad: que nada salga de su equipo.",
        "Una puesta en marcha sencilla, sin conocimientos técnicos.",
        "Visualizaciones claras que se entiendan de un vistazo.",
    ])

    h2(doc, "2.4 Soluciones existentes y comparación")
    para(doc, "A continuación se comparan los enfoques habituales con la propuesta de "
              "Rage Tracker:")
    table_simple(
        doc,
        ["Solución", "Privacidad", "Hardware extra", "Multi-señal", "Coste"],
        [
            ["Sensores biométricos (pulseras)", "Variable", "Sí", "No", "Alto"],
            ["Servicios de análisis facial en la nube", "Baja", "No", "Parcial", "Suscripción"],
            ["Autoinforme manual", "Total", "No", "No", "Gratis"],
            ["Rage Tracker", "Total (local)", "No (webcam/mic)", "Sí (3 señales)", "Gratis"],
        ],
    )
    para(doc, "Rage Tracker es la única opción que combina las tres señales "
              "(cara, gritos e insultos), no requiere hardware adicional más allá de "
              "una webcam y un micrófono corrientes, y mantiene el procesamiento "
              "totalmente local y gratuito.")

    # =====================================================================
    # 3. REQUISITOS
    # =====================================================================
    h1(doc, "3. Requisitos")

    h2(doc, "3.1 Requisitos funcionales")
    para(doc, "Detección de emociones. El sistema captura vídeo de la webcam y "
              "clasifica el rostro en feliz o enfadado con un modelo binario basado en "
              "la presencia de sonrisa. Incorpora votación temporal e histéresis para "
              "evitar parpadeos y falsos positivos, y un sistema de confianza visible "
              "en el HUD.")
    para(doc, "Detección de gritos. El monitor de audio mide el volumen del micrófono "
              "y, cuando supera un umbral configurable durante un tiempo mínimo, "
              "registra un grito. El umbral es ajustable incluso arrastrándolo sobre la "
              "barra VU.")
    para(doc, "Detección de insultos. Un motor de voz a texto en español reconoce el "
              "habla y compara cada palabra con un léxico de insultos. Nunca se muestra "
              "ni se guarda la transcripción.")
    para(doc, "Dashboard de estadísticas. Un servidor local sirve un dashboard web que "
              "muestra estadísticas globales, ranking de juegos y registro de sesiones.")
    para(doc, "Gestión de sesiones. Cada partida se guarda como una fila en un CSV con "
              "todas sus métricas; el sistema permite añadir juegos y consultar el "
              "histórico.")

    h2(doc, "3.2 Requisitos no funcionales")
    bullets(doc, [
        "Rendimiento en tiempo real: la detección facial y de audio debe procesarse a "
        "ritmo de vídeo (objetivo ~25-30 FPS) sin bloquear la interfaz.",
        "Compatibilidad: debe funcionar en Windows, macOS y Linux, con Python 3.8 o "
        "superior, y degradar con elegancia si falta alguna dependencia opcional.",
        "Privacidad de datos: todo el procesamiento es local; no se graba vídeo y "
        "ninguna transcripción de voz se persiste.",
        "Usabilidad: la aplicación arranca con doble clic, ofrece una GUI con tema "
        "oscuro y cae a tkinter estándar si no hay CustomTkinter.",
    ])

    # =====================================================================
    # 4. SOLUCIÓN PROPUESTA
    # =====================================================================
    h1(doc, "4. Solución propuesta")

    h2(doc, "4.1 Descripción general")
    para(doc, "La solución es una aplicación de escritorio en Python organizada en "
              "módulos independientes que se orquestan en torno a una sesión de juego. "
              "El usuario abre el launcher, elige un juego y los sensores que quiere "
              "activar, y lanza la sesión. Mientras juega, los detectores trabajan en "
              "paralelo y un HUD superpuesto muestra el estado en vivo. Al terminar, el "
              "resumen se guarda y puede consultarse en el dashboard.")

    h2(doc, "4.2 Arquitectura del sistema")
    para(doc, "La arquitectura sigue un patrón de orquestador + sensores + "
              "persistencia + visualización. El launcher (GUI) lanza cada sesión como "
              "un subproceso (main.py --session) para que el bucle de eventos de Tkinter "
              "y la ventana de OpenCV no compitan por el hilo principal, algo "
              "especialmente delicado en macOS. El orquestador session_runner combina "
              "los sensores elegidos, recoge sus resúmenes y los persiste vía "
              "data_manager. El dashboard es un servidor HTTP local que lee los CSV y "
              "expone una API JSON consumida por el front-end.")
    photo(doc, "Diagrama general de la arquitectura del sistema")

    h2(doc, "4.3 Módulos principales")
    para(doc, "Detección facial (src/camera.py). Implementa EmotionDetector: abre la "
              "cámara, ejecuta el pipeline de detección de rostro y emoción, pinta el "
              "HUD y produce el resumen de la sesión.")
    para(doc, "Detección de audio (src/audio_monitor.py). Implementa AudioMonitor: "
              "captura bloques de audio, convierte RMS a dBFS y a porcentaje, suaviza la "
              "lectura y detecta gritos. Usa sounddevice como backend principal y "
              "PyAudio como alternativa.")
    para(doc, "Detección de insultos (src/insult_detector.py). Implementa "
              "InsultDetector y un stemmer español propio: abre un stream a 16 kHz, pasa "
              "el audio por Vosk, stemiza cada palabra y la compara con el léxico, con "
              "un debounce de 2 segundos por insulto.")
    para(doc, "Dashboard (web/dashboard_server.py + web/dashboard.html). Servidor HTTP "
              "local con API JSON y front-end con Chart.js.")
    para(doc, "Gestión de datos (src/data_manager.py). Encapsula la lectura y escritura "
              "de los CSV de juegos y sesiones, con migración automática de esquema.")

    # =====================================================================
    # 5. TECNOLOGÍAS UTILIZADAS
    # =====================================================================
    h1(doc, "5. Tecnologías utilizadas")
    para(doc, "Cada tecnología se eligió buscando el equilibrio entre potencia, "
              "facilidad de empaquetado y ausencia de dependencias en la nube.")

    h2(doc, "5.1 Python")
    para(doc, "Lenguaje principal del proyecto (3.8+). Elegido por su ecosistema "
              "maduro en visión por computadora y audio, su rapidez de desarrollo y la "
              "facilidad para empaquetar el resultado como ejecutable. Es además la "
              "tecnología central de la titulación.")

    h2(doc, "5.2 OpenCV")
    para(doc, "Biblioteca de visión por computadora usada para la captura de la "
              "webcam, el preprocesado (CLAHE) y los clasificadores Haar de rostro, "
              "sonrisa y ojos. Se eligió por ser el estándar de facto, funcionar sin "
              "GPU y permitir un pipeline de detección ligero y en tiempo real sin "
              "modelos de deep learning pesados.")

    h2(doc, "5.3 Vosk")
    para(doc, "Motor de reconocimiento de voz (STT) offline. Elegido frente a "
              "alternativas en la nube precisamente porque funciona 100 % local, lo que "
              "es un requisito de privacidad irrenunciable: el audio del micrófono "
              "nunca abandona el equipo. Su modelo en español es ligero (~200 MB) y "
              "suficiente para detectar palabras concretas del léxico.")

    h2(doc, "5.4 SoundDevice")
    para(doc, "Biblioteca de captura de audio en tiempo real. Se eligió como backend "
              "principal por su API limpia basada en callbacks y arrays de NumPy, su "
              "buena compatibilidad multiplataforma y su facilidad de empaquetado. "
              "PyAudio queda como alternativa de respaldo.")

    h2(doc, "5.5 CustomTkinter")
    para(doc, "Framework de interfaz gráfica nativa con tema oscuro moderno. Se eligió "
              "sobre opciones más pesadas (como PyQt) porque, si no está instalado, la "
              "aplicación cae automáticamente a tkinter estándar (incluido en Python), "
              "de modo que la GUI siempre arranca y el empaquetado del .exe se "
              "simplifica.")

    h2(doc, "5.6 HTML / CSS / JavaScript")
    para(doc, "Tecnologías del dashboard web. Se eligieron para desacoplar la "
              "visualización del núcleo en Python: el dashboard se sirve localmente y se "
              "abre en el navegador, lo que permite un diseño rico (estilo "
              "cyberpunk/gaming) y responsive sin complicar la aplicación de escritorio.")

    h2(doc, "5.7 Chart.js")
    para(doc, "Librería JavaScript de gráficos. Elegida por su sencillez, su buen "
              "rendimiento con conjuntos de datos modestos y la calidad visual de sus "
              "gráficos interactivos, ideales para los timelines y rankings del "
              "dashboard.")

    # =====================================================================
    # 6. DISEÑO DEL SISTEMA
    # =====================================================================
    h1(doc, "6. Diseño del sistema")

    h2(doc, "6.1 Arquitectura general")
    para(doc, "El sistema se divide en cuatro capas: captura (cámara y micrófono), "
              "procesamiento (los tres detectores), orquestación y persistencia "
              "(session_runner y data_manager) y presentación (HUD en vivo y dashboard "
              "web). El punto de entrada único es main.py, que según los argumentos "
              "abre la GUI, el menú CLI, una sesión o la calibración.")

    h2(doc, "6.2 Diagrama de módulos (descripción)")
    para(doc, "Descrito textualmente, el flujo de dependencias es el siguiente:")
    bullets(doc, [
        "main.py es el dispatcher: según los flags, llama a launcher, menu, "
        "session_runner o calibration.",
        "launcher.py (GUI) lanza sesiones como subproceso de main.py --session y abre "
        "el dashboard en el navegador.",
        "session_runner.py instancia los sensores pedidos: EmotionDetector (camera.py), "
        "AudioMonitor (audio_monitor.py) y/o InsultDetector (insult_detector.py).",
        "camera.py se apoya en hud.py para el overlay y en calibration.py para el "
        "pipeline de rostro y el perfil de calibración.",
        "Todos los resúmenes desembocan en data_manager.py, que escribe en data/*.csv.",
        "dashboard_server.py lee esos CSV y los expone como JSON a dashboard.html.",
        "paths.py centraliza la resolución de rutas para que todo funcione igual en "
        "desarrollo y empaquetado como .exe.",
    ])
    photo(doc, "Diagrama de módulos y dependencias")

    h2(doc, "6.3 Flujo de funcionamiento")
    numbered(doc, [
        "El usuario abre la aplicación (main.py) y aparece el launcher.",
        "Configura la sesión: elige juego, sensores y ajustes de micrófono.",
        "Al iniciar, se lanza un subproceso que arranca los detectores seleccionados.",
        "Durante la partida, cada detector actualiza su estado y el HUD lo refleja en "
        "vivo (emoción, barra VU, contador de insultos).",
        "El usuario pulsa Q para terminar; cada sensor entrega su resumen.",
        "session_runner fusiona los resúmenes, acopla gritos e insultos al índice de "
        "rage y guarda la sesión en sessions.csv.",
        "El usuario abre el dashboard para revisar las estadísticas acumuladas.",
    ])

    h2(doc, "6.4 Diseño de almacenamiento de datos")
    para(doc, "Se optó por CSV en lugar de una base de datos por ser portable, "
              "legible, sin servidor y fácil de inspeccionar. Hay dos ficheros: "
              "games.csv (catálogo de juegos) y sessions.csv (una fila por sesión). El "
              "esquema de sesiones incluye las métricas de emociones y, de forma "
              "opcional, las de gritos e insultos. La migración de esquema es "
              "automática y atómica: si un CSV antiguo no tiene las columnas nuevas, se "
              "reescribe con un fichero temporal y os.replace para no corromper el "
              "original.")
    para(doc, "Cabecera del esquema canónico de sessions.csv:")
    code_block(doc,
        "game, date, duration_seconds,\n"
        "happy_count, angry_count, neutral_count,\n"
        "happy_percentage, angry_percentage, neutral_percentage,\n"
        "peak_rage_count, happiness_streaks, emotional_trend, total_frames,\n"
        "scream_count, scream_peak_db, scream_total_seconds, mic_device_name,\n"
        "insult_count, insult_peak_count, insult_model_name")

    # =====================================================================
    # 7. IMPLEMENTACIÓN
    # =====================================================================
    h1(doc, "7. Implementación")

    h2(doc, "7.1 Estructura del proyecto")
    code_block(doc,
        "reage-traker/\n"
        "|-- data/                  # Datos persistentes (CSV + perfil de calibración)\n"
        "|   |-- games.csv\n"
        "|   |-- sessions.csv\n"
        "|   |-- insultos.csv\n"
        "|   `-- calibration_profile.json\n"
        "|-- models/vosk-es/        # Modelo Vosk en español (descarga aparte)\n"
        "|-- src/                   # Codigo fuente principal\n"
        "|   |-- audio_monitor.py   # Monitor de microfono / gritos\n"
        "|   |-- calibration.py     # Calibracion automatica de camara\n"
        "|   |-- camera.py          # Detector de emociones faciales\n"
        "|   |-- data_manager.py    # Persistencia en CSV\n"
        "|   |-- hud.py             # HUD sobre el video\n"
        "|   |-- insult_detector.py # Deteccion de insultos (Vosk)\n"
        "|   |-- launcher.py        # GUI nativa (CustomTkinter / tkinter)\n"
        "|   |-- menu.py            # Menu de terminal (CLI)\n"
        "|   |-- paths.py           # Resolucion de rutas (dev / .exe)\n"
        "|   `-- session_runner.py  # Orquestador de sesion multi-sensor\n"
        "|-- utils/config_tool.py   # Herramienta de configuracion\n"
        "|-- web/                   # Dashboard web (server + HTML)\n"
        "|-- tests/                 # Tests unitarios y e2e\n"
        "|-- main.py                # Punto de entrada unificado\n"
        "`-- requirements.txt")

    h2(doc, "7.2 Desarrollo de cada módulo")
    para(doc, "Detección facial. El EmotionDetector usa un pipeline robusto: "
              "preprocesado con CLAHE, detección de rostro con cascada principal y "
              "fallback, selección de la cara más grande, ROI canónica alineada por los "
              "ojos y búsqueda de la sonrisa solo en la región de la boca. La decisión "
              "es binaria con votación temporal:")
    code_block(doc,
        "if smiling:\n"
        "    confidence = int(min(100, 60 + 45 * max(ratio, smile_thr)))\n"
        "    return \"happy\", confidence\n"
        "if eye_count == 0:\n"
        "    # Cara girada/tapada: medicion no fiable, no se cuenta\n"
        "    return \"neutral\", 0\n"
        "confidence = int(min(90, 60 + 30 * (1.0 - ratio)))\n"
        "return \"angry\", confidence")
    photo(doc, "HUD de detección facial en vivo durante una sesión")

    para(doc, "Detección de gritos. Cada bloque de audio se convierte a un nivel 0-100 "
              "y, si supera el umbral durante al menos 0,3 s, se contabiliza un grito:")
    code_block(doc,
        "rms = float(np.sqrt(np.mean(np.square(samples)))) * self.sensitivity\n"
        "db  = _rms_to_db(rms)\n"
        "pct = self._db_to_pct(db)\n"
        "self.level = self.smoothing * self.level + (1 - self.smoothing) * pct\n"
        "if self.level >= self.threshold_pct:\n"
        "    self._pending_above_s += block_seconds\n"
        "    if not self._in_scream and self._pending_above_s >= self.MIN_SCREAM_S:\n"
        "        self._scream_count += 1")

    para(doc, "Detección de insultos. El stemmer español propio reduce cada palabra a "
              "su raíz para que todas las variantes (tonto/tonta/tontos/tontas) colapsen "
              "al mismo radical y casen con el léxico:")
    code_block(doc,
        "SUFFIXES = (\"iendo\", \"ando\", \"amos\", \"ais\",\n"
        "            \"os\", \"as\", \"es\", \"an\", \"en\",\n"
        "            \"ar\", \"er\", \"ir\", \"o\", \"a\")\n\n"
        "for suffix in SpanishStemmer.SUFFIXES:\n"
        "    if word.endswith(suffix):\n"
        "        stem = word[:-len(suffix)]\n"
        "        if len(stem) >= 2:\n"
        "            return stem")
    para(doc, "Por privacidad, el detector solo incrementa contadores; la "
              "transcripción nunca se expone.")

    h2(doc, "7.3 Integración de componentes")
    para(doc, "El EmotionDetector acepta opcionalmente un audio_monitor y un "
              "insult_detector mediante duck-typing: si están presentes, dibuja sus "
              "indicadores en el HUD y, al cerrar la sesión, fusiona sus resúmenes. Así, "
              "una sesión «full» (cámara + gritos + insultos) reutiliza el mismo bucle "
              "de la cámara y simplemente añade los overlays de micrófono e insultos.")

    h2(doc, "7.4 Gestión de sesiones")
    para(doc, "session_runner decide el modo según los sensores: con cámara, ejecuta "
              "EmotionDetector; sin cámara (solo audio), abre una ventana ligera con la "
              "barra VU y los contadores. Tras la sesión, acopla los gritos e insultos "
              "al índice de rage (cada grito suma 1.0 y cada insulto 0.3 «momentos de "
              "enfado») y recalcula los porcentajes, pero solo cuando había cámara, para "
              "no forzar un 100 % de rage artificial en sesiones de solo audio.")
    table_simple(
        doc,
        ["Combinación", "Cámara", "Gritos", "Insultos"],
        [
            ["Solo emociones", "Sí", "No", "No"],
            ["Solo gritos", "No", "Sí", "No"],
            ["Solo insultos", "No", "No", "Sí"],
            ["Gritos + insultos", "No", "Sí", "Sí"],
            ["Emociones + gritos", "Sí", "Sí", "No"],
            ["Emociones + insultos", "Sí", "No", "Sí"],
            ["Full (todos)", "Sí", "Sí", "Sí"],
        ],
    )

    # =====================================================================
    # 8. SEGURIDAD Y PRIVACIDAD
    # =====================================================================
    h1(doc, "8. Seguridad y privacidad")

    h2(doc, "8.1 Procesamiento local")
    para(doc, "Todo el procesamiento ocurre en el equipo del usuario. No hay llamadas "
              "de red, ni telemetría, ni dependencia de servicios en la nube. El modelo "
              "de voz se descarga una sola vez y, a partir de ahí, la aplicación "
              "funciona sin conexión a internet.")

    h2(doc, "8.2 Protección de datos")
    bullets(doc, [
        "No se graba vídeo en ningún momento: los frames se procesan y se descartan.",
        "La detección de insultos no muestra ni almacena la transcripción del habla; "
        "solo se guarda un contador.",
        "Los datos persistidos (CSV) se quedan en la carpeta del usuario y son "
        "auditables y editables por él.",
        "El código es abierto y, por tanto, verificable.",
    ])

    h2(doc, "8.3 Gestión de permisos (cámara y micrófono)")
    para(doc, "La aplicación solicita acceso a la cámara y al micrófono solo cuando se "
              "activan los sensores correspondientes. Si el sistema operativo deniega el "
              "permiso o no hay dispositivo, el módulo afectado se desactiva con un "
              "mensaje de diagnóstico (por ejemplo, la función diagnose() del monitor de "
              "audio) en lugar de bloquear toda la aplicación.")

    h2(doc, "8.4 Limitaciones del sistema")
    bullets(doc, [
        "La detección facial depende de la iluminación y del ángulo: una cara muy "
        "girada o a contraluz reduce la fiabilidad.",
        "La detección de gritos se basa en volumen, por lo que un entorno muy ruidoso "
        "puede generar falsos positivos.",
        "El reconocimiento de insultos está limitado por el modelo Vosk «small» y por "
        "el léxico definido; puede no captar pronunciaciones poco claras.",
        "El modelo emocional es binario (feliz/enfadado); no distingue emociones más "
        "matizadas.",
    ])

    # =====================================================================
    # 9. PRUEBAS
    # =====================================================================
    h1(doc, "9. Pruebas")

    h2(doc, "9.1 Pruebas funcionales")
    para(doc, "Se verificó que cada combinación de sensores arranca, mide y guarda "
              "correctamente la sesión, y que la migración de esquema de sessions.csv "
              "funciona sobre ficheros de versiones anteriores. El proyecto incluye una "
              "batería de tests automatizados con pytest, organizada por marcadores "
              "unit / integration / e2e:")
    code_block(doc,
        "tests/\n"
        "|-- test_csv_migration.py     # Migracion no destructiva del esquema\n"
        "|-- test_fold_insults.py      # Acoplamiento de insultos al rage\n"
        "|-- test_insult_detector.py   # Deteccion y conteo de insultos\n"
        "|-- test_insult_e2e.py        # Flujo extremo a extremo de insultos\n"
        "`-- test_insult_stemmer.py    # Stemmer espanol")

    h2(doc, "9.2 Pruebas de precisión")
    para(doc, "Para la detección facial se usó la calibración asistida, que captura "
              "fases neutral/feliz/enfadado y realiza una búsqueda en rejilla de "
              "parámetros, midiendo la separabilidad entre clases. El perfil obtenido "
              "registra una métrica de calidad; en el equipo de pruebas la separación "
              "happy/no-happy alcanzó valores en torno a 0,94, indicando una "
              "clasificación fiable de la sonrisa.")
    para(doc, "Para los insultos se probó que el stemmer colapsa correctamente las "
              "variantes de género y número a una misma raíz y que el debounce de 2 s "
              "evita el doble conteo de una misma palabra.")
    photo(doc, "Proceso de calibración de cámara (fases neutral / feliz / enfadado)")

    h2(doc, "9.3 Pruebas de rendimiento")
    para(doc, "Se comprobó que la sesión con cámara se mantiene fluida: el HUD muestra "
              "los FPS en vivo y, en hardware de gama media sin GPU dedicada, el "
              "pipeline se sostiene en valores adecuados para tiempo real. Los "
              "detectores de audio corren en hilos/callbacks independientes, de modo que "
              "no penalizan la tasa de refresco del vídeo.")

    h2(doc, "9.4 Resultados obtenidos")
    bullets(doc, [
        "Las tres detecciones funcionan de forma simultánea sin degradar la fluidez.",
        "La persistencia y la migración de datos son fiables y no destructivas.",
        "El dashboard refleja correctamente las métricas acumuladas por juego y sesión.",
        "La aplicación arranca y degrada con elegancia ante la ausencia de "
        "dependencias opcionales o de dispositivos.",
    ])

    # =====================================================================
    # 10. VIABILIDAD Y MEJORAS FUTURAS
    # =====================================================================
    h1(doc, "10. Viabilidad y mejoras futuras")

    h2(doc, "10.1 Viabilidad técnica")
    para(doc, "El proyecto es plenamente viable con tecnología libre y sin coste: se "
              "apoya en bibliotecas maduras, no necesita GPU ni servicios de pago y se "
              "empaqueta como ejecutable de Windows autocontenido. La arquitectura "
              "modular facilita su mantenimiento y ampliación.")

    h2(doc, "10.2 Limitaciones actuales")
    bullets(doc, [
        "Repertorio emocional reducido (binario).",
        "Dependencia de buenas condiciones de luz y de un micrófono razonable.",
        "Léxico de insultos y modelo de voz limitados al español y al diccionario "
        "incluido.",
        "Sin sincronización entre dispositivos ni copia en la nube (por diseño).",
    ])

    h2(doc, "10.3 Posibles mejoras")
    bullets(doc, [
        "IA más avanzada: sustituir las cascadas Haar por un modelo de deep learning "
        "para una clasificación emocional más precisa.",
        "Más emociones detectables: tristeza, sorpresa, asco, miedo, etc.",
        "Aplicación móvil: una versión para registrar sesiones desde el teléfono.",
        "Más métricas de análisis: correlación entre franjas horarias y rage, "
        "detección de fatiga, recomendaciones de descanso.",
    ])

    # =====================================================================
    # 11. MANUAL DE USUARIO
    # =====================================================================
    h1(doc, "11. Manual de usuario")

    h2(doc, "11.1 Instalación")
    para(doc, "Opción A — Ejecutable de Windows (recomendado para usuarios). Descarga "
              "el ZIP de la página del proyecto o de las Releases, descomprímelo y abre "
              "RageTracker.exe. No requiere instalar Python ni nada más.")
    photo(doc, "Carpeta descomprimida con RageTracker.exe")
    para(doc, "Opción B — Desde el código fuente. Requiere Python 3.8 o superior:")
    code_block(doc,
        "# 1. Clonar el repositorio\n"
        "git clone https://github.com/Anthony0827/reage-traker.git\n"
        "cd reage-traker\n\n"
        "# 2. Crear y activar el entorno virtual\n"
        "python -m venv venv\n"
        "venv\\Scripts\\activate        # Windows\n"
        "source venv/bin/activate      # Linux / Mac\n\n"
        "# 3. Instalar dependencias\n"
        "pip install -r requirements.txt\n\n"
        "# 4. (Insultos) Descargar el modelo Vosk espanol en models/vosk-es/\n\n"
        "# 5. Ejecutar\n"
        "python main.py")

    h2(doc, "11.2 Configuración")
    para(doc, "Al abrir la aplicación aparece el launcher. Pulsa «CONFIGURAR SESIÓN» "
              "para elegir el juego, activar los sensores (emociones, gritos, insultos) "
              "y ajustar el micrófono. El panel incluye un medidor de volumen en vivo y "
              "una línea de umbral que puedes arrastrar sobre la barra VU. También puedes "
              "calibrar la cámara desde aquí para mejorar la detección facial.")
    photo(doc, "Panel de configuración de sesión con selección de sensores y medidor de micrófono")
    para(doc, "Ajustes rápidos por línea de comandos (modo avanzado):")
    code_block(doc,
        "python main.py --session --game \"CS2\" \\\n"
        "    --sensors emotions scream insults \\\n"
        "    --threshold 80 --sensitivity 1.5")

    h2(doc, "11.3 Uso de la aplicación")
    numbered(doc, [
        "Pulsa «CONFIGURAR SESIÓN» y selecciona un juego (o añade uno nuevo).",
        "Elige los sensores que quieras activar.",
        "Ajusta el micrófono (umbral, sensibilidad y dispositivo).",
        "Pulsa «INICIAR RAGE TRACKER» y juega con normalidad.",
        "Observa el HUD: emoción actual, barra de volumen y contador de insultos.",
        "Pulsa Q para terminar y guardar la sesión (R reinicia los contadores, P "
        "pausa).",
        "Abre el dashboard para revisar tus estadísticas.",
    ])
    photo(doc, "Sesión en curso con el HUD superpuesto sobre la imagen de la webcam")
    para(doc, "Controles durante la sesión:")
    table_simple(
        doc,
        ["Tecla", "Acción"],
        [
            ["Q", "Terminar y guardar la sesión"],
            ["R", "Reiniciar todos los contadores"],
            ["P", "Pausar / reanudar"],
            ["C", "Recalibrar la cámara en caliente"],
        ],
    )

    h2(doc, "11.4 Interpretación de estadísticas")
    para(doc, "El dashboard se abre en el navegador (http://localhost:8000) y se "
              "organiza en tres pestañas: Dashboard (resumen global con tarjetas, "
              "timeline y heatmap), Juegos (ranking y comparativa por juego, incluida la "
              "«boca más sucia») y Registro (listado de todas las sesiones).")
    photo(doc, "Pestaña principal del dashboard con las estadísticas globales")
    para(doc, "Claves de lectura:")
    bullets(doc, [
        "Rage / porcentaje de enfado: cuanto mayor, más frustración acumulada en la "
        "sesión (incluye gritos e insultos cuando había cámara).",
        "Gritos: número de veces que el volumen superó el umbral el tiempo suficiente.",
        "Insultos: número de insultos del léxico detectados (sin transcripción).",
        "Rachas de felicidad: secuencias sostenidas de detecciones «feliz».",
        "Picos de rage: momentos puntuales de enfado de alta confianza.",
    ])
    photo(doc, "Vista de comparación entre juegos en el dashboard")

    # =====================================================================
    # 12. CONCLUSIONES
    # =====================================================================
    h1(doc, "12. Conclusiones")

    h2(doc, "12.1 Objetivos alcanzados")
    para(doc, "Se ha construido un sistema multi-sensor funcional que detecta "
              "emociones, gritos e insultos en tiempo real, los persiste de forma fiable "
              "y los presenta en un dashboard claro, cumpliendo el objetivo principal y "
              "los específicos. La premisa de privacidad total —procesamiento 100 % "
              "local y sin guardar vídeo ni transcripciones— se ha respetado en todo el "
              "diseño.")

    h2(doc, "12.2 Aprendizajes obtenidos")
    bullets(doc, [
        "Diseño de un pipeline de visión por computadora robusto sin recurrir a deep "
        "learning, cuidando la normalización (CLAHE, ROI canónica, alineación por "
        "ojos).",
        "Captura y análisis de audio en tiempo real con backends intercambiables y "
        "tolerancia a fallos de hardware.",
        "Integración de un motor de voz a texto offline respetando la privacidad.",
        "Orquestación de procesos y de hilos para que Tkinter y OpenCV convivan sin "
        "bloqueos.",
        "Diseño de persistencia evolutiva con migración de esquema no destructiva.",
    ])

    h2(doc, "12.3 Valor del proyecto")
    para(doc, "Rage Tracker demuestra que es posible ofrecer una herramienta útil de "
              "autoconocimiento emocional para jugadores sin sacrificar la privacidad ni "
              "exigir hardware especializado. Su arquitectura modular y su empaquetado "
              "sencillo lo hacen fácil de mantener y de ampliar, y sienta una base "
              "sólida sobre la que incorporar las mejoras futuras planteadas.")

    # ---- Guardar ----
    out_dir = Path(__file__).resolve().parents[1] / "documentacion"
    out_dir.mkdir(exist_ok=True)
    out = out_dir / "Rage_Tracker_Memoria.docx"
    doc.save(str(out))
    print(f"Documento generado: {out}")


if __name__ == "__main__":
    build()
